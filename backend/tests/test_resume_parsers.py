from io import BytesIO

import pytest
import pymupdf
from docx import Document

from app.resume_parsers import (
    DocxResumeParser,
    InvalidDocumentError,
    NoExtractableTextError,
    PdfResumeParser,
    ResumeParserRegistry,
)


@pytest.fixture
def pdf_bytes():
    document = pymupdf.open()
    first_page = document.new_page()
    first_page.insert_textbox(
        pymupdf.Rect(72, 72, 520, 180),
        "Project experience: built a retrieval augmented generation agent.",
    )
    second_page = document.new_page()
    second_page.insert_textbox(
        pymupdf.Rect(72, 72, 520, 180),
        "Technology stack: Python, FastAPI, vector database, and observability.",
    )
    try:
        return document.tobytes()
    finally:
        document.close()


@pytest.fixture
def empty_pdf_bytes():
    document = pymupdf.open()
    document.new_page()
    try:
        return document.tobytes()
    finally:
        document.close()


@pytest.fixture
def docx_bytes():
    document = Document()
    document.add_paragraph(
        "工作经历：负责企业知识库 Agent 后端开发，设计检索链路并维护线上稳定性。"
    )
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "FastAPI 和 RAG"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_pdf_parser_extracts_text_and_page_count(pdf_bytes):
    parsed = PdfResumeParser().parse(pdf_bytes, "resume.pdf", "application/pdf")

    assert parsed.source_type == "pdf"
    assert parsed.unit_count == 2
    assert "Project experience" in parsed.text
    assert "Technology stack" in parsed.text


def test_docx_parser_extracts_paragraphs_and_tables(docx_bytes):
    parsed = DocxResumeParser().parse(
        docx_bytes,
        "resume.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert parsed.source_type == "docx"
    assert "工作经历" in parsed.text
    assert "Python" in parsed.text
    assert "FastAPI 和 RAG" in parsed.text
    assert parsed.unit_count >= 2


def test_pdf_parser_rejects_empty_text(empty_pdf_bytes):
    with pytest.raises(NoExtractableTextError) as caught:
        PdfResumeParser().parse(empty_pdf_bytes, "resume.pdf", "application/pdf")

    assert caught.value.code == "no_extractable_text"


def test_parsers_reject_corrupted_documents():
    with pytest.raises(InvalidDocumentError):
        PdfResumeParser().parse(b"not a pdf", "resume.pdf", "application/pdf")

    with pytest.raises(InvalidDocumentError):
        DocxResumeParser().parse(b"not a docx", "resume.docx", None)


def test_docx_parser_rejects_insufficient_text():
    document = Document()
    document.add_paragraph("too short")
    output = BytesIO()
    document.save(output)

    with pytest.raises(NoExtractableTextError):
        DocxResumeParser().parse(output.getvalue(), "resume.docx", None)


def test_registry_selects_parser_by_extension(pdf_bytes, docx_bytes):
    assert isinstance(
        ResumeParserRegistry.for_file("resume.pdf", "application/pdf"),
        PdfResumeParser,
    )
    assert isinstance(
        ResumeParserRegistry.for_file("resume.docx", None),
        DocxResumeParser,
    )
