from io import BytesIO

import pymupdf
from docx import Document
from sqlalchemy import select

from app.models import Resume, ResumeSource, User


PROJECT = {
    "project_name": "知识库 Agent",
    "background_goal": "提升检索效率",
    "tech_stack": "Python、FastAPI、向量数据库",
    "responsibilities": "负责后端和评估",
    "core_solution": "混合检索和重排",
    "engineering_challenges": "召回质量",
    "failure_improvements": "增加监控和降级",
    "quantified_results": "P95 降低 20%",
}


def make_pdf_bytes():
    document = pymupdf.open()
    page = document.new_page()
    page.insert_textbox(
        pymupdf.Rect(72, 72, 520, 180),
        "Project experience: built a retrieval augmented generation agent for enterprise search.",
    )
    try:
        return document.tobytes()
    finally:
        document.close()


def make_docx_bytes():
    document = Document()
    document.add_paragraph(
        "工作经历：负责企业知识库 Agent 后端开发，设计检索链路并维护线上稳定性。"
    )
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "FastAPI"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_profile_and_session_creation_have_fixed_shape(client):
    profile_response = client.post(
        "/api/profile",
        json={
            "resume_text": "我负责过一个 RAG Agent 项目。",
            "project": {
                "project_name": "知识库 Agent",
                "background_goal": "提升检索效率",
                "tech_stack": "Python、FastAPI、向量数据库",
                "responsibilities": "负责后端和评估",
                "core_solution": "混合检索和重排",
                "engineering_challenges": "召回质量",
                "failure_improvements": "增加监控和降级",
                "quantified_results": "P95 降低 20%",
            },
        },
    )

    assert profile_response.status_code == 200
    profile = profile_response.json()
    assert profile["user_id"] == "local-user"
    assert profile["project_version"] == 1
    assert len(profile["resume_text_hash"]) == 64

    session_response = client.post(
        "/api/sessions", json={"profile_id": profile["profile_id"]}
    )

    assert session_response.status_code == 200
    session = session_response.json()
    assert len(session["questions"]) == 8
    assert [question["category"] for question in session["questions"]].count("project") == 3
    assert [question["category"] for question in session["questions"]].count("agent") == 3
    assert [question["category"] for question in session["questions"]].count("reliability") == 2
    assert [question["is_anchor"] for question in session["questions"][:3]] == [True, True, True]


def test_parse_pdf_persists_resume_source(client):
    response = client.post(
        "/api/resumes/parse",
        files={"file": ("resume.pdf", make_pdf_bytes(), "application/pdf")},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["source_type"] == "pdf"
    assert body["unit_count"] == 1
    assert body["extracted_text"]
    assert body["character_count"] == len(body["extracted_text"])

    with client.app.state.session_factory() as db:
        assert len(list(db.scalars(select(Resume)))) == 1
        sources = list(db.scalars(select(ResumeSource)))
        assert len(sources) == 1
        assert sources[0].source_type == "pdf"
        assert (client.app.state.upload_root / sources[0].stored_path).exists()


def test_parse_docx_returns_extracted_text(client):
    response = client.post(
        "/api/resumes/parse",
        files={
            "file": (
                "resume.docx",
                make_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert response.json()["source_type"] == "docx"
    assert "工作经历" in response.json()["extracted_text"]
    assert "Python" in response.json()["extracted_text"]


def test_parse_returns_stable_file_errors(client):
    unsupported = client.post(
        "/api/resumes/parse",
        files={"file": ("resume.txt", b"plain text", "text/plain")},
    )
    invalid_signature = client.post(
        "/api/resumes/parse",
        files={"file": ("resume.pdf", b"not a pdf", "application/pdf")},
    )
    too_large = client.post(
        "/api/resumes/parse",
        files={"file": ("resume.pdf", b"x" * (10 * 1024 * 1024 + 1), "application/pdf")},
    )

    assert unsupported.status_code == 415
    assert unsupported.json()["code"] == "unsupported_file_type"
    assert invalid_signature.status_code == 422
    assert invalid_signature.json()["code"] == "invalid_file_signature"
    assert too_large.status_code == 413
    assert too_large.json()["code"] == "file_too_large"


def test_profile_reuses_parsed_resume_and_saves_final_text(client):
    parsed = client.post(
        "/api/resumes/parse",
        files={"file": ("resume.pdf", make_pdf_bytes(), "application/pdf")},
    ).json()
    final_text = parsed["extracted_text"] + "\n用户确认补充：负责线上故障复盘。"

    response = client.post(
        "/api/profile",
        json={"resume_id": parsed["resume_id"], "resume_text": final_text, "project": PROJECT},
    )

    assert response.status_code == 200
    with client.app.state.session_factory() as db:
        resumes = list(db.scalars(select(Resume)))
        assert len(resumes) == 1
        assert resumes[0].resume_text == final_text


def test_profile_rejects_resume_owned_by_another_user(client):
    with client.app.state.session_factory() as db:
        db.add(User(id="other-user"))
        db.add(
            Resume(
                id="other-resume",
                user_id="other-user",
                resume_text="这是一份用于权限测试的简历文本。",
                text_hash="0" * 64,
            )
        )
        db.commit()

    response = client.post(
        "/api/profile",
        json={"resume_id": "other-resume", "resume_text": "修改后的文本", "project": PROJECT},
    )

    assert response.status_code == 409
    assert response.json()["code"] == "resume_owner_conflict"


def test_session_read_returns_current_question_and_progress(client, session_context):
    session_id, questions = session_context

    response = client.get(f"/api/sessions/{session_id}")

    assert response.status_code == 200
    view = response.json()
    assert view["session_id"] == session_id
    assert view["status"] == "in_progress"
    assert view["current_question"]["id"] == questions[0]["id"]
    assert view["progress"] == {"completed": 0, "total": 8}


def test_duplicate_submission_returns_one_answer_and_one_observation(
    client, session_context
):
    session_id, questions = session_context
    payload = {
        "question_id": questions[0]["id"],
        "client_submission_id": "stable-1",
        "status": "submitted",
        "answer_text": "我负责业务目标和检索链路，结果通过评估集验证并降低了延迟。",
    }

    first = client.post(f"/api/sessions/{session_id}/answers", json=payload)
    second = client.post(f"/api/sessions/{session_id}/answers", json=payload)

    assert first.status_code == 200
    assert second.status_code == 200
    assert first.json()["answer"]["id"] == second.json()["answer"]["id"]
    assert first.json()["observation"]["id"] == second.json()["observation"]["id"]


def test_same_submission_id_with_different_payload_returns_409(client, session_context):
    session_id, questions = session_context
    first_payload = {
        "question_id": questions[0]["id"],
        "client_submission_id": "stable-conflict",
        "status": "submitted",
        "answer_text": "第一版回答。",
    }
    second_payload = {**first_payload, "answer_text": "不同内容的第二版回答。"}

    assert client.post(f"/api/sessions/{session_id}/answers", json=first_payload).status_code == 200
    conflict = client.post(f"/api/sessions/{session_id}/answers", json=second_payload)

    assert conflict.status_code == 409


def test_unknown_and_skipped_have_distinct_evaluation_behavior(client, session_context):
    session_id, questions = session_context
    unknown = client.post(
        f"/api/sessions/{session_id}/answers",
        json={
            "question_id": questions[0]["id"],
            "client_submission_id": "unknown-1",
            "status": "explicit_unknown",
            "answer_text": "不知道",
        },
    )
    skipped = client.post(
        f"/api/sessions/{session_id}/answers",
        json={
            "question_id": questions[1]["id"],
            "client_submission_id": "skip-1",
            "status": "skipped",
            "answer_text": "",
        },
    )

    assert unknown.status_code == 200
    assert unknown.json()["observation"]["level"] == 0
    assert skipped.status_code == 200
    assert skipped.json()["observation"] is None


def test_blank_submitted_answer_is_rejected(client, session_context):
    session_id, questions = session_context
    response = client.post(
        f"/api/sessions/{session_id}/answers",
        json={
            "question_id": questions[0]["id"],
            "client_submission_id": "blank-1",
            "status": "submitted",
            "answer_text": "   ",
        },
    )

    assert response.status_code == 422


def test_report_aggregates_without_uncalibrated_score(client, session_context):
    session_id, questions = session_context
    for index, question in enumerate(questions[:2]):
        response = client.post(
            f"/api/sessions/{session_id}/answers",
            json={
                "question_id": question["id"],
                "client_submission_id": f"report-{index}",
                "status": "submitted",
                "answer_text": "我会结合机制、边界、监控和评估集分析问题并验证结果。",
            },
        )
        assert response.status_code == 200

    report_response = client.get(f"/api/sessions/{session_id}/report")

    assert report_response.status_code == 200
    report = report_response.json()
    assert report["completion"] == {"completed": 2, "total": 8}
    assert report["anchor_coverage"] == {"answered": 2, "total": 3}
    assert 0 < report["coverage"] < 1
    assert report["valid_evidence_count"] == 2
    assert "score_100" not in report
